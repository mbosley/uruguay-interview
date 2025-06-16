"""
Document processing module for interview ingestion.
"""
from pathlib import Path
from typing import Dict, Any, Optional
import re
from dataclasses import dataclass
from datetime import datetime
import logging

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx not installed. DOCX support disabled.")

logger = logging.getLogger(__name__)


@dataclass
class InterviewDocument:
    """Represents a processed interview document."""
    id: str
    date: str
    time: str
    location: str
    department: Optional[str]
    participant_count: int
    text: str
    metadata: Dict[str, Any]
    file_path: str


class DocumentProcessor:
    """Processes interview documents from various formats."""
    
    # Location mappings for Uruguay departments
    URUGUAY_DEPARTMENTS = {
        'montevideo', 'canelones', 'maldonado', 'rocha', 'treinta y tres',
        'cerro largo', 'rivera', 'artigas', 'salto', 'paysandú', 'río negro',
        'soriano', 'colonia', 'san josé', 'flores', 'florida', 'lavalleja',
        'durazno', 'tacuarembó'
    }
    
    def process_interview(self, file_path: Path) -> InterviewDocument:
        """Extract text and metadata from interview document."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() == '.docx':
            if docx is None:
                raise ImportError("python-docx is required for DOCX files")
            return self._process_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def _process_docx(self, file_path: Path) -> InterviewDocument:
        """Process DOCX file."""
        doc = docx.Document(str(file_path))
        
        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        text = '\n'.join(full_text)
        
        # Extract metadata
        metadata = self._extract_metadata(file_path.name, text)
        
        return InterviewDocument(
            id=metadata['id'],
            date=metadata['date'],
            time=metadata['time'],
            location=metadata['location'],
            department=metadata['department'],
            participant_count=metadata['participant_count'],
            text=text,
            metadata=metadata,
            file_path=str(file_path)
        )
    
    def _process_txt(self, file_path: Path) -> InterviewDocument:
        """Process TXT file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Extract metadata
        metadata = self._extract_metadata(file_path.name, text)
        
        return InterviewDocument(
            id=metadata['id'],
            date=metadata['date'],
            time=metadata['time'],
            location=metadata['location'],
            department=metadata['department'],
            participant_count=metadata['participant_count'],
            text=text,
            metadata=metadata,
            file_path=str(file_path)
        )
    
    def _extract_metadata(self, filename: str, text: str) -> Dict[str, Any]:
        """Extract metadata from filename and text content."""
        metadata = {}
        
        # Extract from filename
        # Expected format: "20250528_0900_058.txt" or "T - 20250528 0900 58.docx"
        filename_clean = filename.replace('.txt', '').replace('.docx', '')
        
        # Try different filename patterns
        # Pattern 1: YYYYMMDD_HHMM_ID
        match = re.match(r'(\d{8})_(\d{4})_(\d+)', filename_clean)
        if match:
            date_str, time_str, id_str = match.groups()
            metadata['date'] = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
            metadata['time'] = f"{time_str[:2]}:{time_str[2:]}"
            metadata['id'] = id_str
        else:
            # Pattern 2: T - YYYYMMDD HHMM ID
            match = re.match(r'T\s*-\s*(\d{8})\s+(\d{4})\s+(\d+)', filename_clean)
            if match:
                date_str, time_str, id_str = match.groups()
                metadata['date'] = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                metadata['time'] = f"{time_str[:2]}:{time_str[2:]}"
                metadata['id'] = id_str
            else:
                # Fallback
                metadata['date'] = datetime.now().strftime('%Y-%m-%d')
                metadata['time'] = '00:00'
                metadata['id'] = filename_clean
        
        # Extract location and participant info from text
        metadata['location'] = self._detect_location(text)
        metadata['department'] = self._detect_department(text)
        metadata['participant_count'] = self._count_participants(text)
        
        # Additional metadata
        metadata['filename'] = filename
        metadata['text_length'] = len(text)
        metadata['word_count'] = len(text.split())
        
        return metadata
    
    def _detect_location(self, text: str) -> str:
        """Detect interview location from text."""
        # Look for location patterns in first 500 chars
        text_start = text[:500].lower()
        
        # Common location indicators
        location_patterns = [
            r'lugar:\s*([^\n]+)',
            r'ubicación:\s*([^\n]+)',
            r'localidad:\s*([^\n]+)',
            r'ciudad:\s*([^\n]+)',
            r'barrio:\s*([^\n]+)'
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text_start)
            if match:
                return match.group(1).strip().title()
        
        # Check for department names
        for dept in self.URUGUAY_DEPARTMENTS:
            if dept in text_start:
                return dept.title()
        
        return "Unknown"
    
    def _detect_department(self, text: str) -> Optional[str]:
        """Detect Uruguay department from text."""
        text_lower = text[:1000].lower()
        
        for dept in self.URUGUAY_DEPARTMENTS:
            if dept in text_lower:
                return dept.title()
        
        return None
    
    def _count_participants(self, text: str) -> int:
        """Count number of participants in interview."""
        # Look for speaker indicators
        speaker_patterns = [
            r'^[A-Z]{2,3}:',  # Initials like "AM:", "JP:"
            r'^\w+:',  # Names like "Juan:"
            r'Entrevistado\s*\d*:',
            r'Participante\s*\d*:'
        ]
        
        speakers = set()
        lines = text.split('\n')
        
        for line in lines[:100]:  # Check first 100 lines
            for pattern in speaker_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    speaker = match.group(0).rstrip(':')
                    # Exclude common interviewer markers
                    if speaker.lower() not in ['entrevistador', 'e', 'ent']:
                        speakers.add(speaker)
        
        # If no speakers found, assume 1 participant
        return max(len(speakers), 1)


if __name__ == "__main__":
    # Test with a sample file
    processor = DocumentProcessor()
    
    # Example usage
    sample_file = Path("data/processed/interviews_txt/20250528_0900_058.txt")
    if sample_file.exists():
        interview = processor.process_interview(sample_file)
        print(f"Processed interview {interview.id}")
        print(f"Date: {interview.date} {interview.time}")
        print(f"Location: {interview.location}")
        print(f"Participants: {interview.participant_count}")
        print(f"Text length: {len(interview.text)} characters")